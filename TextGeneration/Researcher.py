from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate
from BaseApp import BaseApp
import streamlit as st
from langchain_core.tools import Tool
from langchain_google_community import GoogleSearchAPIWrapper
import PyPDF2
import docx
import pandas as pd
from Model_Manager import ModelManager
from Config import PRIMARY_MODEL, SECONDARY_MODEL, Images_MODEL
from PIL import Image
import base64
import io
from langchain_openai import ChatOpenAI
from langchain_anthropic import ChatAnthropic

class Researcher(BaseApp):
    def __init__(self, model_manager, app_name="The Researcher 🔬📚", app_slogan="Your AI-Powered Research Assistant! 🤖✨"):
        super().__init__(app_name)
        self.app_slogan = app_slogan
        self.model_manager = model_manager

        if "chat_history" not in st.session_state:
            st.session_state.chat_history = []
        if "uploaded_doc" not in st.session_state:
            st.session_state.uploaded_doc = None
        if "user_selections" not in st.session_state:
            st.session_state.user_selections = []
        if "original_image" not in st.session_state:
            st.session_state.original_image = None
        if "image_summary" not in st.session_state:
            st.session_state.image_summary = None
        if "COST" not in st.session_state:
            st.session_state.COST = 0
        if "research_papers" not in st.session_state:
            st.session_state.research_papers = False

    def extract_text_from_pdf(self, pdf_file):
        """Extracts text from a PDF file."""
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        return "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])

    def extract_text_from_docx(self, docx_file):
        """Extracts text from a DOCX file."""
        doc = docx.Document(docx_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)

    def encode_image(self, image_file):
        """Encodes an image as a base64 string."""
        img = Image.open(image_file).convert("RGB").resize((1024, 1024))
        buffered = io.BytesIO()
        img.save(buffered, format="JPEG", quality=75)
        return base64.b64encode(buffered.getvalue()).decode("utf-8")

    def handle_image_query(self, user_input, image_content):
        """Handles image-based queries using GPT-4o-mini (fallback: Claude 3.5 Sonnet) and tracks cost."""
        from Model_Manager import PRICING  # Import pricing dictionary

        try:
            model_name = PRIMARY_MODEL["model_name"]
            model = ChatOpenAI(model=model_name, temperature=PRIMARY_MODEL["temperature"])
            used_model = "GPT-4o-mini (Primary)"
        except Exception as e:
            print(f"Primary Model Failed: {e}... Switching to Claude 3.5 Sonnet")
            model_name = Images_MODEL["model_name"]
            model = ChatAnthropic(model=model_name, temperature=Images_MODEL["temperature"])
            used_model = "Claude 3.5 Sonnet (Fallback)"

        content = [
            {"type": "text", "text": "Summarize this image before answering the query."},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_content}" }},
            {"type": "text", "text": user_input}
        ]

        try:
            response = model.invoke([HumanMessage(content=content)])
            summary_response = response.content
            st.session_state.image_summary = summary_response

            input_tokens = response.usage_metadata.get("input_tokens", 0)
            output_tokens = response.usage_metadata.get("output_tokens", 0)

            cost_summary = (input_tokens * PRICING[model_name]["input"]) + \
                        (output_tokens * PRICING[model_name]["output"])
            st.session_state.COST += cost_summary

            query_prompt = ChatPromptTemplate.from_template("""
                {summary}
                Answer the following query in details if needed (don't state that you are summarizing but rather you know the info): {query}
            """)

            chain = query_prompt | model
            final_response = chain.invoke({"summary": summary_response, "query": user_input})

            input_tokens = final_response.usage_metadata.get("input_tokens", 0)
            output_tokens = final_response.usage_metadata.get("output_tokens", 0)

            cost_query = (input_tokens * PRICING[model_name]["input"]) + \
                        (output_tokens * PRICING[model_name]["output"])
            st.session_state.COST += cost_query

            print(f"### 🔍 Model Used: `{used_model}`")
            print(f"💰 **Total Cost:** ${st.session_state.COST:.6f}")

            print(model_name)
            print(st.session_state.COST)
            st.session_state.COST = 0 

            return final_response.content

        except Exception as e:
            print(f"Primary model failed: {e}. Switching to Claude 3.5 Sonnet...")
            secondary_model = ChatAnthropic(model=Images_MODEL["model_name"], temperature=Images_MODEL["temperature"])
            used_model = "Claude 3.5 Sonnet (Fallback)"

            fallback_response = secondary_model.invoke([HumanMessage(content=content)])

            input_tokens = fallback_response.usage_metadata.get("input_tokens", 0)
            output_tokens = fallback_response.usage_metadata.get("output_tokens", 0)

            cost_fallback = (input_tokens * PRICING[Images_MODEL["model_name"]]["input"]) + \
                            (output_tokens * PRICING[Images_MODEL["model_name"]]["output"])
            st.session_state.COST += cost_fallback

            
            print(st.session_state.COST)
            st.session_state.COST = 0  

            return fallback_response.content



    def generate_prompt(self):
        """Creates a structured research prompt."""
        return ChatPromptTemplate.from_template("""
            You are ResearchGenie, an advanced AI-powered research assistant. Base responses on:

            - Web search results: {web_results}
            - Uploaded document: {uploaded_doc}
            - Suggest research papers if {research_papers}
            If no sources are available, use your own knowledge.

            User Question: {user_input}
        """)

    def display_side(self):
        """Displays sidebar options for user settings and file uploads."""
        with st.sidebar:
            st.header("🔎 Research Settings")
            st.session_state.web_search_enabled = st.toggle("Enable Web Search")

            uploaded_file = st.file_uploader("📂 Upload Document or Image", type=["pdf", "docx", "csv", "jpg", "jpeg", "png"])
            research_papers = st.toggle("Research Papers")
            st.session_state.research_papers = research_papers
            if uploaded_file:
                file_type = uploaded_file.name.split(".")[-1].lower()
                content = None

                if file_type in ["jpg", "jpeg", "png"]:
                    content = self.encode_image(uploaded_file)
                    st.session_state.original_image = uploaded_file.getvalue()
                elif file_type == "pdf":
                    content = self.extract_text_from_pdf(uploaded_file)
                elif file_type == "docx":
                    content = self.extract_text_from_docx(uploaded_file)
                elif file_type == "csv":
                    df = pd.read_csv(uploaded_file, encoding="latin1")
                    content = df.to_string()

                st.session_state.uploaded_doc = {"type": file_type, "name": uploaded_file.name, "content": content}

            if st.button("🧹 Clear Upload"):
                st.session_state.uploaded_doc = None
                st.session_state.original_image = None
                st.rerun()

    def handle_input(self):
        """Handles user input for research queries."""
        user_input = st.chat_input("Type your research question here...")
        if user_input:
            with st.chat_message("User", avatar="😀"):
                st.markdown(user_input)
            st.session_state.chat_history.append(HumanMessage(content=user_input))

            final_response = ""
            if st.session_state.uploaded_doc and st.session_state.uploaded_doc["type"] in ["jpg", "jpeg", "png"]:
                try:
                    final_response = self.handle_image_query(user_input, st.session_state.uploaded_doc["content"])
                except Exception as e:
                    final_response = f"Error processing image: {str(e)}"
            else:
                prompt_template = self.generate_prompt()
                response, st.session_state.COST = self.model_manager.generate(
                    prompt_template,
                    {
                        "web_results": "Web search disabled",
                        "uploaded_doc": st.session_state.uploaded_doc["content"] if st.session_state.uploaded_doc else "None",
                        "user_input": user_input,
                        "research_papers": st.session_state.research_papers
                    }
                )
                print(st.session_state.COST)
                final_response = response

            super().display_ai_response(final_response)
            st.session_state.chat_history.append(AIMessage(content=final_response))

def main():
    model_manager = ModelManager(PRIMARY_MODEL, SECONDARY_MODEL)
    app = Researcher(model_manager)
    app.welcome_screen(app.app_slogan)
    app.display_side()
    app.handle_input()

if __name__ == "__main__":
    main()
